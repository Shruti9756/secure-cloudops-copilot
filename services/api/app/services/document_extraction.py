"""Safely extract bounded plain text from supported document upload bytes."""

from io import BytesIO
from zipfile import BadZipFile, is_zipfile

import pymupdf
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

MAX_EXTRACTED_TEXT_CHARACTERS = 1_000_000


def extract_document_text(*, suffix: str, content_bytes: bytes) -> str:
    """Return safe plain text for one supported uploaded document format."""
    normalized_suffix = suffix.lower()

    if normalized_suffix in {".md", ".txt"}:
        return _extract_utf8_text(content_bytes)

    if normalized_suffix == ".pdf":
        return _extract_pdf_text(content_bytes)

    if normalized_suffix == ".docx":
        return _extract_docx_text(content_bytes)

    raise ValueError("Only .md, .txt, .pdf, and .docx uploads are supported")


def _extract_utf8_text(content_bytes: bytes) -> str:
    """Decode the existing plain-text formats without changing their meaning."""
    try:
        content = content_bytes.decode("utf-8")
    except UnicodeDecodeError as error:
        raise ValueError("Text uploads must use UTF-8 encoding") from error

    return _combine_extracted_parts(
        [content],
        empty_message="Uploaded document must contain readable text",
    )


def _extract_pdf_text(content_bytes: bytes) -> str:
    """Extract digital PDF text while retaining page-level provenance markers."""
    if not content_bytes.startswith(b"%PDF-"):
        raise ValueError("Uploaded .pdf file is not a valid PDF")

    try:
        with pymupdf.open(stream=content_bytes, filetype="pdf") as pdf_document:
            parts = []

            for page_number, page in enumerate(pdf_document, start=1):
                page_text = page.get_text("text", sort=True).strip()

                if page_text:
                    # Keep page provenance so future answers can identify their source.
                    parts.append(f"[PDF page {page_number}]\n{page_text}")
    except (pymupdf.FileDataError, RuntimeError) as error:
        raise ValueError("Uploaded .pdf file could not be read safely") from error

    return _combine_extracted_parts(
        parts,
        empty_message=("PDF does not contain extractable text; scanned PDFs are not supported"),
    )


def _extract_docx_text(content_bytes: bytes) -> str:
    """Extract readable DOCX headings, paragraphs, and tables without executing code."""
    if not is_zipfile(BytesIO(content_bytes)):
        raise ValueError("Uploaded .docx file is not a valid DOCX archive")

    try:
        document = Document(BytesIO(content_bytes))
    except (BadZipFile, PackageNotFoundError, ValueError) as error:
        raise ValueError("Uploaded .docx file could not be read safely") from error

    parts = []

    for paragraph in document.paragraphs:
        formatted_paragraph = _format_docx_paragraph(paragraph.text, paragraph.style.name)

        if formatted_paragraph:
            parts.append(formatted_paragraph)

    for table in document.tables:
        formatted_table = _format_docx_table(table)

        if formatted_table:
            parts.append(formatted_table)

    return _combine_extracted_parts(
        parts,
        empty_message="DOCX does not contain extractable text",
    )


def _format_docx_paragraph(text: str, style_name: str) -> str:
    """Preserve Word headings as Markdown-like headings for useful RAG chunks."""
    normalized_text = text.strip()

    if not normalized_text:
        return ""

    if style_name.startswith("Heading"):
        level_text = style_name.removeprefix("Heading").strip()
        heading_level = int(level_text) if level_text.isdigit() else 1
        bounded_level = min(max(heading_level, 1), 6)
        return f"{'#' * bounded_level} {normalized_text}"

    return normalized_text


def _format_docx_table(table: object) -> str:
    """Convert a DOCX table into readable plain text without interpreting its content."""
    rows = []

    for row in table.rows:
        cell_values = [" ".join(cell.text.split()) for cell in row.cells]

        if any(cell_values):
            rows.append(f"| {' | '.join(cell_values)} |")

    if not rows:
        return ""

    return "[DOCX table]\n" + "\n".join(rows)


def _combine_extracted_parts(
    parts: list[str],
    *,
    empty_message: str,
) -> str:
    """Join non-empty text while enforcing a limit before RAG processing starts."""
    combined_parts = []
    character_count = 0

    for part in parts:
        normalized_part = part.strip()

        if not normalized_part:
            continue

        separator_length = 2 if combined_parts else 0
        next_character_count = character_count + separator_length + len(normalized_part)

        if next_character_count > MAX_EXTRACTED_TEXT_CHARACTERS:
            raise ValueError("Extracted document text exceeds the 1,000,000 character safety limit")

        combined_parts.append(normalized_part)
        character_count = next_character_count

    if not combined_parts:
        raise ValueError(empty_message)

    return "\n\n".join(combined_parts)
