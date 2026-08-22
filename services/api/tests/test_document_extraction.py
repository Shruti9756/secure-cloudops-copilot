"""Tests for safe PDF and DOCX text extraction."""

from io import BytesIO

import pymupdf
import pytest
from docx import Document

from app.services.document_extraction import extract_document_text


def make_pdf_bytes(text: str | None = None) -> bytes:
    """Build an in-memory PDF so tests do not need real uploaded files."""
    pdf_document = pymupdf.open()

    try:
        page = pdf_document.new_page()

        if text is not None:
            page.insert_text((72, 72), text)

        return pdf_document.tobytes()
    finally:
        pdf_document.close()


def make_docx_bytes() -> bytes:
    """Build an in-memory DOCX containing the content we expect to preserve."""
    document = Document()
    document.add_heading("Checkout Investigation", level=1)
    document.add_paragraph("Inspect PostgreSQL connections after the deployment.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Signal"
    table.cell(0, 1).text = "Action"
    table.cell(1, 0).text = "Redis eviction"
    table.cell(1, 1).text = "Inspect eviction policy"

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_extract_document_text_preserves_pdf_page_provenance() -> None:
    """PDF text remains readable and identifies the page it came from."""
    content = extract_document_text(
        suffix=".pdf",
        content_bytes=make_pdf_bytes("Checkout latency increased after deployment."),
    )

    assert "[PDF page 1]" in content
    assert "Checkout latency increased after deployment." in content


def test_extract_document_text_rejects_scanned_or_empty_pdf() -> None:
    """An image-only or blank PDF cannot silently become empty RAG knowledge."""
    with pytest.raises(
        ValueError,
        match="PDF does not contain extractable text",
    ):
        extract_document_text(
            suffix=".pdf",
            content_bytes=make_pdf_bytes(),
        )


def test_extract_document_text_extracts_docx_headings_paragraphs_and_tables() -> None:
    """DOCX content becomes readable RAG text without executing document content."""
    content = extract_document_text(
        suffix=".docx",
        content_bytes=make_docx_bytes(),
    )

    assert "# Checkout Investigation" in content
    assert "Inspect PostgreSQL connections after the deployment." in content
    assert "[DOCX table]" in content
    assert "| Redis eviction | Inspect eviction policy |" in content


def test_extract_document_text_rejects_invalid_docx_archive() -> None:
    """A file renamed to .docx must not be treated as a Word document."""
    with pytest.raises(ValueError, match="not a valid DOCX archive"):
        extract_document_text(
            suffix=".docx",
            content_bytes=b"This is not a DOCX file.",
        )


def test_extract_document_text_keeps_utf8_markdown_supported() -> None:
    """Existing Markdown uploads still use the same text extraction path."""
    content = extract_document_text(
        suffix=".md",
        content_bytes=b"# Redis Investigation\n\nInspect eviction policy.",
    )

    assert content == "# Redis Investigation\n\nInspect eviction policy."
